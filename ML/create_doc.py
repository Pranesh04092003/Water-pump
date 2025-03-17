from docx import Document
from docx.shared import Inches
import markdown
import os

def create_word_doc():
    # Create a new Word document
    doc = Document()
    
    # Read markdown content
    with open('documentation.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Split content by sections (##)
    sections = md_content.split('## ')
    
    # Add title
    doc.add_heading(sections[0].strip(), 0)
    
    # Process each section
    for section in sections[1:]:
        # Add section heading
        lines = section.split('\n')
        doc.add_heading(lines[0], 1)
        
        # Add content
        content = '\n'.join(lines[1:])
        doc.add_paragraph(content)
        
        # Add relevant plots based on section
        if "Results and Analysis" in lines[0]:
            doc.add_heading("Visualization Results", 2)
            
            # Add Vibration Analysis plots
            doc.add_heading("Vibration Analysis", 3)
            doc.add_picture('plots/vibration/vibration_distribution.png', width=Inches(6))
            doc.add_paragraph("Figure 1: Vibration Distribution by Motor Status")
            
            doc.add_picture('plots/vibration/confusion_matrix.png', width=Inches(6))
            doc.add_paragraph("Figure 2: Vibration Model Confusion Matrix")
            
            # Add Cooling Analysis plots
            doc.add_heading("Cooling Efficiency Analysis", 3)
            doc.add_picture('plots/cooling/cooling_metrics.png', width=Inches(6))
            doc.add_paragraph("Figure 3: Cooling Efficiency Metrics")
            
            doc.add_picture('plots/cooling/confusion_matrix.png', width=Inches(6))
            doc.add_paragraph("Figure 4: Cooling Model Confusion Matrix")
            
            # Add Feature Analysis plots
            doc.add_heading("Feature Analysis", 3)
            doc.add_picture('plots/feature_importance.png', width=Inches(6))
            doc.add_paragraph("Figure 5: Feature Importance for Both Models")
            
            doc.add_picture('plots/cooling/correlation_matrix.png', width=Inches(6))
            doc.add_paragraph("Figure 6: Cooling Features Correlation Matrix")
            
            # Add Time Series Analysis
            doc.add_heading("Time Series Analysis", 3)
            doc.add_picture('plots/time_series.png', width=Inches(6))
            doc.add_paragraph("Figure 7: Time Series Analysis of Vibration and Cooling")
            
            # Add ROC Curves
            doc.add_heading("Model Performance Curves", 3)
            doc.add_picture('plots/roc_curves.png', width=Inches(6))
            doc.add_paragraph("Figure 8: ROC Curves for Both Models")
    
    # Save the document
    doc.save('Motor_Cooling_ML_Documentation.docx')
    print("✅ Word document created successfully with all plots!")

if __name__ == "__main__":
    create_word_doc()

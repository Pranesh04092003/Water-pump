from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import os

def create_conference_paper():
    # Create a new Document
    doc = Document()
    
    # Read the markdown content
    with open('conference_paper.docx', 'r') as f:
        content = f.read()
    
    # Split content into sections
    sections = content.split('\n## ')
    
    # Process title
    title = sections[0].split('\n')[0].replace('# ', '')
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process each section
    for section in sections[1:]:
        lines = section.split('\n')
        section_title = lines[0]
        doc.add_heading(section_title, 1)
        
        # Process content
        content = '\n'.join(lines[1:])
        
        # Check for image placeholders
        if '[Insert' in content:
            # Add text before image
            text_before = content.split('[Insert')[0]
            if text_before.strip():
                doc.add_paragraph(text_before)
            
            # Add image
            image_name = content.split('[Insert ')[1].split(']')[0]
            if os.path.exists(f'plots/{image_name}'):
                doc.add_picture(f'plots/{image_name}', width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add text after image
            text_after = content.split(']')[1]
            if text_after.strip():
                doc.add_paragraph(text_after)
        else:
            # Add regular text
            doc.add_paragraph(content)
    
    # Save the document
    doc.save('Advanced_Motor_Pattern_Learning.docx')
    print("✅ Conference paper generated successfully!")

if __name__ == "__main__":
    create_conference_paper()
